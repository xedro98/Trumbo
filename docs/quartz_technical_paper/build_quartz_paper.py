from __future__ import annotations

from pathlib import Path
from typing import Iterable

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
OUT = ROOT / "Quartz_Technical_Paper.docx"
FIG_DIR = ROOT / "figures"

SERIF = "Times New Roman"
SANS = "Arial"

NAVY = "171717"
BLUE = "334155"
TEAL = "2F5D62"
GREEN = "3F6656"
GOLD = "7A5A28"
RED = "7A3B3B"
INK = "202020"
MUTED = "5D6168"
LINE = "BFC5CE"
PALE_BLUE = "EEF2F6"
PALE_TEAL = "EEF6F5"
PALE_GREEN = "F0F6F2"
PALE_GOLD = "FAF4E8"
PALE_RED = "F7EEEE"
PALE_GRAY = "F6F6F6"
WHITE = "FFFFFF"


def rgb(value: str) -> RGBColor:
    return RGBColor.from_string(value)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top: int = 100, start: int = 120, bottom: int = 100, end: int = 120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color: str = LINE, size: str = "6") -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = qn(f"w:{edge}")
        node = borders.find(tag)
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def set_picture_alt(inline_shape, title: str, description: str) -> None:
    doc_pr = inline_shape._inline.docPr
    doc_pr.set("title", title)
    doc_pr.set("descr", description)


def set_table_widths(table, widths_in: list[float]) -> None:
    table.autofit = False
    total_dxa = int(round(sum(widths_in) * 1440))
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total_dxa))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for col, width in zip(grid.gridCol_lst, widths_in):
        col.set(qn("w:w"), str(int(round(width * 1440))))
    for row in table.rows:
        for cell, width in zip(row.cells, widths_in):
            cell.width = Inches(width)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(round(width * 1440))))
            tc_w.set(qn("w:type"), "dxa")


def set_font(
    run,
    size: float | None = None,
    bold: bool | None = None,
    color: str | None = None,
    italic: bool | None = None,
    font: str = SERIF,
) -> None:
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:ascii"), font)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), font)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = rgb(color)
    if italic is not None:
        run.italic = italic


def set_para(paragraph, before: float = 0, after: float = 8, line: float = 1.25, keep: bool = False) -> None:
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    if keep:
        fmt.keep_with_next = True


def add_text(
    doc: Document,
    text: str,
    *,
    style: str = "Normal",
    before: float = 0,
    after: float = 7,
    line: float = 1.16,
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
) -> None:
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    set_para(p, before=before, after=after, line=line, keep=style.startswith("Heading"))
    set_font(p.add_run(text))


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_paragraph(style=f"Heading {level}")
    p.add_run(text)
    set_para(p, before={1: 15, 2: 10, 3: 7}[level], after={1: 6, 2: 4, 3: 3}[level], line=1.0, keep=True)


def add_caption(doc: Document, label: str, title: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.left_indent = Inches(0.18)
    p.paragraph_format.right_indent = Inches(0.18)
    set_para(p, before=1, after=10, line=1.0)
    r = p.add_run(f"{label}. ")
    set_font(r, size=9.2, bold=True, color=INK)
    r = p.add_run(title)
    set_font(r, size=9.2, color=INK)


def add_figure(doc: Document, path: Path, label: str, title: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para(p, before=7, after=0, line=1.0, keep=True)
    inline_shape = p.add_run().add_picture(str(path), width=Inches(6.2))
    set_picture_alt(inline_shape, f"{label}: {title}", f"Logical architecture diagram for {title.lower()}.")
    add_caption(doc, label, title)


def add_callout(doc: Document, title: str, body: str, fill: str = PALE_BLUE) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_widths(table, [6.5])
    set_table_borders(table, color=LINE, size="4")
    set_repeat_table_header(table.rows[0])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_margins(cell, top=120, bottom=120, start=165, end=165)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    set_para(p, after=2, line=1.0)
    r = p.add_run(title)
    set_font(r, size=9.8, bold=True, color=INK)
    p = cell.add_paragraph()
    set_para(p, after=0, line=1.1)
    set_font(p.add_run(body), size=9.8, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_widths(table, widths)
    set_table_borders(table, color="AEB4BC", size="4")
    header = table.rows[0]
    set_repeat_table_header(header)
    for cell, text in zip(header.cells, headers):
        set_cell_shading(cell, "F2F2F2")
        set_cell_margins(cell, top=80, bottom=80, start=95, end=95)
        p = cell.paragraphs[0]
        set_para(p, after=0, line=1.0)
        set_font(p.add_run(text), size=8.9, bold=True, color=INK)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for row_values in rows:
        cells = table.add_row().cells
        for cell, text in zip(cells, row_values):
            set_cell_margins(cell, top=80, bottom=80, start=95, end=95)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            p = cell.paragraphs[0]
            set_para(p, after=0, line=1.08)
            set_font(p.add_run(text), size=8.7, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def add_reference(doc: Document, number: int, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    set_para(p, after=3, line=1.05)
    set_font(p.add_run(f"[{number}] {text}"), size=9.0, color=INK)


def draw_rounded(draw: ImageDraw.ImageDraw, box: tuple[int, int, int, int], fill: str, outline: str = LINE, radius: int = 18, width: int = 3) -> None:
    draw.rounded_rectangle(box, radius=radius, fill=f"#{fill}", outline=f"#{outline}", width=width)


def diagram_fonts():
    candidates = [
        Path("C:/Windows/Fonts/arialbd.ttf"),
        Path("C:/Windows/Fonts/Arial.ttf"),
    ]
    bold = next((p for p in candidates if p.exists() and "bd" in p.name.lower()), candidates[0])
    regular = next((p for p in candidates if p.exists() and "bd" not in p.name.lower()), candidates[-1])
    return (
        ImageFont.truetype(str(bold), 30),
        ImageFont.truetype(str(regular), 22),
        ImageFont.truetype(str(bold), 20),
        ImageFont.truetype(str(regular), 18),
    )


def center_text(draw, box, text, font, color=f"#{INK}", line_gap: int = 4):
    max_width = box[2] - box[0] - 24
    words = text.split()
    lines: list[str] = []
    line = ""
    for word in words:
        trial = f"{line} {word}".strip()
        if draw.textbbox((0, 0), trial, font=font)[2] <= max_width or not line:
            line = trial
        else:
            lines.append(line)
            line = word
    if line:
        lines.append(line)
    heights = [draw.textbbox((0, 0), x, font=font)[3] for x in lines]
    total = sum(heights) + line_gap * (len(lines) - 1)
    y = box[1] + ((box[3] - box[1]) - total) // 2
    for line, h in zip(lines, heights):
        w = draw.textbbox((0, 0), line, font=font)[2]
        draw.text((box[0] + ((box[2] - box[0]) - w) // 2, y), line, font=font, fill=color)
        y += h + line_gap


def arrow(draw, start, end, color=BLUE, width=6, label: str | None = None, font=None):
    draw.line([start, end], fill=f"#{color}", width=width)
    ex, ey = end
    sx, sy = start
    dx, dy = ex - sx, ey - sy
    length = max((dx * dx + dy * dy) ** 0.5, 1)
    ux, uy = dx / length, dy / length
    px, py = -uy, ux
    tip = (ex, ey)
    left = (int(ex - 18 * ux + 10 * px), int(ey - 18 * uy + 10 * py))
    right = (int(ex - 18 * ux - 10 * px), int(ey - 18 * uy - 10 * py))
    draw.polygon([tip, left, right], fill=f"#{color}")
    if label and font:
        mx, my = (sx + ex) // 2, (sy + ey) // 2
        bbox = draw.textbbox((0, 0), label, font=font)
        draw.rounded_rectangle((mx - bbox[2] // 2 - 6, my - bbox[3] // 2 - 4, mx + bbox[2] // 2 + 6, my + bbox[3] // 2 + 4), radius=5, fill="#FFFFFF")
        draw.text((mx - bbox[2] // 2, my - bbox[3] // 2), label, font=font, fill=f"#{MUTED}")


def make_canvas(title: str, subtitle: str = ""):
    img = Image.new("RGB", (1800, 850), f"#{WHITE}")
    draw = ImageDraw.Draw(img)
    title_font, body_font, bold_small, small = diagram_fonts()
    draw.text((65, 36), title, font=title_font, fill=f"#{NAVY}")
    if subtitle:
        draw.text((65, 79), subtitle, font=small, fill=f"#{MUTED}")
    draw.line((65, 115, 1735, 115), fill=f"#{LINE}", width=3)
    return img, draw, body_font, bold_small, small


def save(img: Image.Image, name: str) -> Path:
    FIG_DIR.mkdir(parents=True, exist_ok=True)
    path = FIG_DIR / name
    img.save(path, quality=95)
    return path


def fig_system_overview() -> Path:
    img, draw, body, bold, small = make_canvas("Quartz system overview", "A closed-weight foundation model system for multimodal reasoning and agent execution")
    columns = [
        (75, 235, 320, 610, PALE_BLUE, "Input fabric", "Text, code, documents, images, tool state"),
        (430, 190, 775, 655, PALE_TEAL, "Quartz core", "Long-context sequence backbone\nSparse expert substrate\nAttention residual path"),
        (910, 190, 1255, 655, PALE_GOLD, "Inference fabric", "Task decomposition\nAdaptive reasoning\nVerifier-guided selection"),
        (1390, 235, 1725, 610, PALE_GREEN, "Response fabric", "Grounded answer\nStructured output\nTool actions and streams"),
    ]
    for x1, y1, x2, y2, fill, head, body_text in columns:
        draw_rounded(draw, (x1, y1, x2, y2), fill)
        center_text(draw, (x1 + 18, y1 + 25, x2 - 18, y1 + 95), head, bold, f"#{NAVY}")
        draw.line((x1 + 28, y1 + 120, x2 - 28, y1 + 120), fill=f"#{LINE}", width=2)
        center_text(draw, (x1 + 26, y1 + 145, x2 - 26, y2 - 28), body_text, body, f"#{INK}")
    arrow(draw, (320, 423), (430, 423), TEAL, label="encode", font=small)
    arrow(draw, (775, 423), (910, 423), GOLD, label="allocate", font=small)
    arrow(draw, (1255, 423), (1390, 423), GREEN, label="synthesize", font=small)
    draw_rounded(draw, (460, 715, 1290, 800), "F8FAFC", LINE, radius=14)
    center_text(draw, (480, 725, 1270, 790), "Persistent operating principles: private weights | bounded tool authority | evidence-aware verification | mode-specific compute budgets", small, f"#{MUTED}")
    return save(img, "figure_01_system_overview.png")


def fig_sequence_core() -> Path:
    img, draw, body, bold, small = make_canvas("Hybrid long-context sequence core", "Two complementary pathways preserve local precision while scaling contextual reach")
    input_box = (80, 315, 310, 505)
    draw_rounded(draw, input_box, PALE_BLUE)
    center_text(draw, input_box, "Token and multimodal embeddings", body)
    upper = (475, 175, 850, 340)
    lower = (475, 510, 850, 675)
    fuse = (1030, 315, 1325, 505)
    out = (1485, 315, 1720, 505)
    for box, fill, label in [
        (upper, PALE_TEAL, "Selective attention path\nHigh-resolution local and retrieval-sensitive interactions"),
        (lower, PALE_GOLD, "Gated linear memory path\nLinear-growth recurrent state for long-range continuity"),
        (fuse, PALE_GREEN, "Residual fusion\nLearned mixing, normalization, and depth-stable transport"),
        (out, PALE_BLUE, "Contextual states"),
    ]:
        draw_rounded(draw, box, fill)
        center_text(draw, box, label, body)
    arrow(draw, (310, 380), (475, 260), TEAL, label="detail", font=small)
    arrow(draw, (310, 450), (475, 590), GOLD, label="memory", font=small)
    arrow(draw, (850, 260), (1030, 370), TEAL)
    arrow(draw, (850, 590), (1030, 450), GOLD)
    arrow(draw, (1325, 410), (1485, 410), GREEN, label="decode", font=small)
    draw_rounded(draw, (410, 730, 1430, 805), "F8FAFC", LINE, radius=14)
    center_text(draw, (435, 740, 1405, 795), "Design intent: reserve quadratic attention for information-dense neighborhoods; use a gated memory pathway to carry stable state across long spans; retain an explicit residual path so useful attention signals remain easy to recover at depth.", small, f"#{MUTED}")
    return save(img, "figure_02_hybrid_sequence_core.png")


def fig_expert_core() -> Path:
    img, draw, body, bold, small = make_canvas("Stable sparse expert execution", "Conditional computation concentrates capacity where a token sequence needs it")
    boxes = [
        (80, 290, 300, 530, PALE_BLUE, "Shared trunk", "Normalized hidden states"),
        (420, 250, 690, 570, PALE_GOLD, "Expert selector", "Scores experts\nBalances load\nApplies capacity guardrails"),
        (845, 165, 1130, 305, PALE_TEAL, "Expert set A", "Code and formal reasoning"),
        (845, 355, 1130, 495, PALE_GREEN, "Expert set B", "Knowledge and language"),
        (845, 545, 1130, 685, PALE_RED, "Expert set C", "Tool and action planning"),
        (1305, 290, 1575, 530, PALE_BLUE, "Gated merge", "Weighted outputs\nResidual carry\nNext block"),
    ]
    for x1, y1, x2, y2, fill, head, text in boxes:
        draw_rounded(draw, (x1, y1, x2, y2), fill)
        center_text(draw, (x1 + 15, y1 + 20, x2 - 15, y1 + 80), head, bold, f"#{NAVY}")
        center_text(draw, (x1 + 20, y1 + 92, x2 - 20, y2 - 20), text, body)
    arrow(draw, (300, 410), (420, 410), BLUE)
    arrow(draw, (690, 365), (845, 235), TEAL, label="top-k", font=small)
    arrow(draw, (690, 410), (845, 425), GREEN, label="top-k", font=small)
    arrow(draw, (690, 455), (845, 615), RED, label="top-k", font=small)
    for start in [(1130, 235), (1130, 425), (1130, 615)]:
        arrow(draw, start, (1305, 410), BLUE)
    draw_rounded(draw, (250, 745, 1580, 810), "F8FAFC", LINE, radius=14)
    center_text(draw, (275, 753, 1555, 802), "The selector is part of the model substrate: it activates internal specialist circuits for each token while shared pathways preserve common language and instruction behavior.", small, f"#{MUTED}")
    return save(img, "figure_03_sparse_expert_core.png")


def fig_memory_fabric() -> Path:
    img, draw, body, bold, small = make_canvas("Long-context memory fabric", "Stable prefixes, segment summaries, and recurrent state reduce repeated work without weakening attribution")
    blocks = [
        (75, 245, 315, 585, PALE_BLUE, "Request assembly", "System context\nConversation\nFiles and media\nTool observations"),
        (430, 175, 745, 355, PALE_TEAL, "Stable-prefix cache", "Hash-aligned reusable prefix states"),
        (430, 475, 745, 655, PALE_GOLD, "Segment index", "Anchors, summaries, source spans"),
        (885, 245, 1200, 585, PALE_GREEN, "Context composer", "Selects attention windows\nRestores cached state\nTags provenance"),
        (1340, 245, 1725, 585, PALE_BLUE, "Quartz core", "Hybrid sequence states\nSparse experts\nReasoning workspace"),
    ]
    for x1, y1, x2, y2, fill, head, text in blocks:
        draw_rounded(draw, (x1, y1, x2, y2), fill)
        center_text(draw, (x1 + 15, y1 + 20, x2 - 15, y1 + 82), head, bold, f"#{NAVY}")
        center_text(draw, (x1 + 20, y1 + 95, x2 - 20, y2 - 20), text, body)
    arrow(draw, (315, 345), (430, 265), TEAL, label="stable prefix", font=small)
    arrow(draw, (315, 485), (430, 565), GOLD, label="segments", font=small)
    arrow(draw, (745, 265), (885, 345), TEAL)
    arrow(draw, (745, 565), (885, 485), GOLD)
    arrow(draw, (1200, 415), (1340, 415), GREEN, label="composed context", font=small)
    return save(img, "figure_04_long_context_memory.png")


def fig_adaptive_inference() -> Path:
    img, draw, body, bold, small = make_canvas("Adaptive inference and verification", "Compute expands only when uncertainty, task depth, or consequence warrants it")
    nodes = [
        (80, 315, 300, 505, PALE_BLUE, "Task signal", "Prompt, context, constraints, prior outcome"),
        (415, 315, 665, 505, PALE_GOLD, "Difficulty estimator", "Complexity, ambiguity, tool need, risk"),
        (805, 165, 1125, 305, PALE_TEAL, "Direct response", "Fast synthesis\nLow-risk tasks"),
        (805, 365, 1125, 505, PALE_GREEN, "Reasoning loop", "Plan, derive, critique, refine"),
        (805, 565, 1125, 705, PALE_RED, "Tool-grounded loop", "Act, observe, update, retry"),
        (1290, 315, 1550, 505, PALE_BLUE, "Verifier", "Check constraints, evidence, result consistency"),
        (1650, 315, 1765, 505, PALE_TEAL, "Answer", "Verified result"),
    ]
    for x1, y1, x2, y2, fill, head, text in nodes:
        draw_rounded(draw, (x1, y1, x2, y2), fill)
        center_text(draw, (x1 + 12, y1 + 15, x2 - 12, y1 + 72), head, bold, f"#{NAVY}")
        center_text(draw, (x1 + 16, y1 + 82, x2 - 16, y2 - 16), text, body)
    arrow(draw, (300, 410), (415, 410), BLUE)
    arrow(draw, (665, 360), (805, 235), TEAL, label="simple", font=small)
    arrow(draw, (665, 410), (805, 435), GREEN, label="complex", font=small)
    arrow(draw, (665, 460), (805, 635), RED, label="interactive", font=small)
    for pos in [(1125, 235), (1125, 435), (1125, 635)]:
        arrow(draw, pos, (1290, 410), BLUE)
    arrow(draw, (1550, 410), (1650, 410), TEAL)
    return save(img, "figure_05_adaptive_inference.png")


def fig_agent_loop() -> Path:
    img, draw, body, bold, small = make_canvas("Tool-grounded execution loop", "Planning, constrained action, observation, and verification form one bounded control cycle")
    nodes = [
        (100, 300, 350, 510, PALE_BLUE, "Goal state", "User objective, permissions, success condition"),
        (500, 150, 800, 330, PALE_TEAL, "Plan", "Subgoals, dependencies, tool eligibility"),
        (1010, 150, 1310, 330, PALE_GOLD, "Act", "Schema-valid tool call with bounded authority"),
        (1450, 300, 1710, 510, PALE_GREEN, "Observe", "Tool result, environment state, error signals"),
        (1010, 575, 1310, 735, PALE_RED, "Verify", "Goal progress, evidence, safety and stop rule"),
        (500, 575, 800, 735, PALE_BLUE, "Synthesize", "Final answer or revised plan"),
    ]
    for x1, y1, x2, y2, fill, head, text in nodes:
        draw_rounded(draw, (x1, y1, x2, y2), fill)
        center_text(draw, (x1 + 15, y1 + 18, x2 - 15, y1 + 75), head, bold, f"#{NAVY}")
        center_text(draw, (x1 + 18, y1 + 85, x2 - 18, y2 - 18), text, body)
    arrow(draw, (350, 350), (500, 240), TEAL)
    arrow(draw, (800, 240), (1010, 240), GOLD)
    arrow(draw, (1310, 300), (1450, 385), GREEN)
    arrow(draw, (1570, 510), (1310, 655), RED)
    arrow(draw, (1010, 655), (800, 655), BLUE)
    arrow(draw, (500, 655), (225, 510), BLUE, label="continue or finalize", font=small)
    return save(img, "figure_06_agent_loop.png")


def fig_output_contract() -> Path:
    img, draw, body, bold, small = make_canvas("Reliable output contract", "Reasoning workspaces and user-visible responses are separated by verification and schema constraints")
    steps = [
        (85, 265, 335, 560, PALE_GOLD, "Internal workspace", "Hypotheses\nCandidate plans\nIntermediate checks"),
        (465, 265, 715, 560, PALE_TEAL, "Evidence gate", "Source attribution\nTool-result binding\nConstraint checks"),
        (845, 265, 1095, 560, PALE_GREEN, "Decoder policy", "Natural language\nStrict JSON Schema\nTool-call schema\nPartial continuation"),
        (1225, 265, 1475, 560, PALE_BLUE, "Stream channels", "Progress stream\nFinal answer stream\nStructured events"),
        (1605, 265, 1745, 560, PALE_TEAL, "Client", "People and software"),
    ]
    for x1, y1, x2, y2, fill, head, text in steps:
        draw_rounded(draw, (x1, y1, x2, y2), fill)
        center_text(draw, (x1 + 14, y1 + 18, x2 - 14, y1 + 78), head, bold, f"#{NAVY}")
        center_text(draw, (x1 + 18, y1 + 92, x2 - 18, y2 - 18), text, body)
    for a, b, color in [((335, 412), (465, 412), TEAL), ((715, 412), (845, 412), GREEN), ((1095, 412), (1225, 412), BLUE), ((1475, 412), (1605, 412), TEAL)]:
        arrow(draw, a, b, color)
    return save(img, "figure_07_output_contract.png")


def fig_training_lifecycle() -> Path:
    img, draw, body, bold, small = make_canvas("Training and alignment lifecycle", "A staged program joins broad competence with controllability, tool reliability, and safety evidence")
    boxes = [
        (85, 245, 330, 560, PALE_BLUE, "Data curation", "Quality filtering\nDeduplication\nContamination control\nMultimodal alignment"),
        (430, 245, 675, 560, PALE_TEAL, "Foundation training", "Long-context curricula\nSparse-expert balance\nSequence stability"),
        (775, 245, 1020, 560, PALE_GOLD, "Capability shaping", "Reasoning traces\nCode and tool tasks\nStructured generation"),
        (1120, 245, 1365, 560, PALE_GREEN, "Alignment", "Preference learning\nProcess signals\nSafety red teaming"),
        (1465, 245, 1710, 560, PALE_RED, "Release gates", "Regression suites\nCapability-risk review\nOperational monitoring"),
    ]
    for x1, y1, x2, y2, fill, head, text in boxes:
        draw_rounded(draw, (x1, y1, x2, y2), fill)
        center_text(draw, (x1 + 12, y1 + 18, x2 - 12, y1 + 78), head, bold, f"#{NAVY}")
        center_text(draw, (x1 + 16, y1 + 92, x2 - 16, y2 - 18), text, body)
    for x in [330, 675, 1020, 1365]:
        arrow(draw, (x, 402), (x + 100, 402), BLUE)
    return save(img, "figure_08_training_lifecycle.png")


def fig_evaluation() -> Path:
    img, draw, body, bold, small = make_canvas("Evaluation and release evidence", "Quartz is evaluated as an integrated model system rather than as a single-turn text generator")
    center = (720, 295, 1080, 555)
    draw_rounded(draw, center, PALE_BLUE)
    center_text(draw, center, "Quartz evaluation harness\nVersioned tasks, hidden holdouts, audited failure traces", body)
    sectors = [
        ((90, 160, 470, 330), PALE_TEAL, "Reasoning", "Mathematics, logical consistency, multi-step planning"),
        ((90, 535, 470, 705), PALE_GREEN, "Engineering", "Repository tasks, tests, debugging, structured edits"),
        ((1330, 160, 1710, 330), PALE_GOLD, "Long context", "Retrieval fidelity, multi-document synthesis, cache behavior"),
        ((1330, 535, 1710, 705), PALE_RED, "Agent safety", "Permission boundaries, tool error recovery, stop conditions"),
    ]
    for box, fill, head, text in sectors:
        draw_rounded(draw, box, fill)
        center_text(draw, (box[0] + 18, box[1] + 16, box[2] - 18, box[1] + 65), head, bold, f"#{NAVY}")
        center_text(draw, (box[0] + 18, box[1] + 78, box[2] - 18, box[3] - 16), text, body)
    for a, b, color in [((470, 245), (720, 355), TEAL), ((470, 620), (720, 495), GREEN), ((1330, 245), (1080, 355), GOLD), ((1330, 620), (1080, 495), RED)]:
        arrow(draw, a, b, color)
    return save(img, "figure_09_evaluation.png")


def build_doc() -> None:
    figures = [
        fig_system_overview(), fig_sequence_core(), fig_expert_core(), fig_memory_fabric(),
        fig_adaptive_inference(), fig_agent_loop(), fig_output_contract(), fig_training_lifecycle(), fig_evaluation(),
    ]

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.9)
    section.bottom_margin = Inches(0.85)
    section.left_margin = Inches(0.95)
    section.right_margin = Inches(0.95)
    section.header_distance = Inches(0.38)
    section.footer_distance = Inches(0.36)
    section.different_first_page_header_footer = True

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = SERIF
    normal._element.rPr.rFonts.set(qn("w:ascii"), SERIF)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), SERIF)
    normal.font.size = Pt(10.6)
    normal.font.color.rgb = rgb(INK)
    for level, size, color in [(1, 13.6, INK), (2, 11.5, INK), (3, 10.8, INK)]:
        style = styles[f"Heading {level}"]
        style.font.name = SERIF
        style._element.rPr.rFonts.set(qn("w:ascii"), SERIF)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), SERIF)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(color)
    if "Caption" not in styles:
        styles.add_style("Caption", WD_STYLE_TYPE.PARAGRAPH)

    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_para(hp, after=0, line=1.0)
    set_font(hp.add_run("Quartz Technical Research Paper"), size=8.6, color=MUTED)
    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para(fp, after=0, line=1.0)
    set_font(fp.add_run("Trumbo Research | July 2026"), size=8.4, color=MUTED)

    # Manuscript title block
    p = doc.add_paragraph()
    set_para(p, before=7, after=5, line=1.0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(p.add_run("Quartz: A Private-Weights, Long-Context Foundation Model for Reasoning and Agentic Work"), size=19, color=INK, bold=True)
    p = doc.add_paragraph()
    set_para(p, after=5, line=1.0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(p.add_run("Technical architecture, adaptive inference, tool-grounded execution, and reliability design"), size=10.8, color=MUTED, italic=True)
    p = doc.add_paragraph()
    set_para(p, after=2, line=1.0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(p.add_run("Shubhankar Kahali, Kai Keskitalo, and Zuzanna Kowalczyk"), size=10.6, color=INK)
    p = doc.add_paragraph()
    set_para(p, after=2, line=1.0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(p.add_run("Trumbo Research, Trumbo"), size=10.0, color=INK, italic=True)
    p = doc.add_paragraph()
    set_para(p, after=8, line=1.0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(p.add_run("Technical Research Paper | July 2026"), size=9.6, color=MUTED)
    rule = doc.add_paragraph()
    set_para(rule, after=8, line=1.0)
    p_pr = rule._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "808080")
    borders.append(bottom)
    p_pr.append(borders)

    # Abstract and scope
    add_heading(doc, "Abstract", 2)
    add_text(doc, "Quartz is Trumbo's closed-source foundation-model system for long-context understanding, multimodal reasoning, and reliable agentic work. This paper defines the architectural thesis behind Quartz: a private sparse-expert core couples a hybrid long-context sequence pathway with adaptive inference, verification, structured generation, and bounded tool use. The design treats quality as a system property. A capable decoder alone is insufficient when a task requires codebase-scale context, high-consequence reasoning, tool execution, or machine-readable outputs. Quartz therefore organizes inference around five interacting layers: input and context assembly, a long-context model backbone, conditional expert computation, adaptive reasoning and action loops, and a reliability contract at the output boundary.")
    add_text(doc, "The paper does not claim a released parameter count or benchmark ranking. Instead, it states the implementation-oriented principles that a frontier private model system should satisfy: preserve evidence across long contexts; apply specialist capacity without turning product behavior into a collection of disconnected systems; expand inference compute in proportion to uncertainty; bind tool actions to explicit schemas and permissions; and emit answers that can be validated by both machines and humans. The resulting architecture is designed to support direct chat, software engineering, document intelligence, structured automation, and long-running agent workflows under one coherent Quartz identity.")
    add_heading(doc, "Keywords", 2)
    add_text(doc, "foundation models; sparse experts; long-context modeling; adaptive inference; process verification; tool use; structured generation; multimodal reasoning; private weights", after=6)
    add_callout(doc, "Scope of Disclosure", "This manuscript describes Quartz design targets, operating assumptions, and evaluation methodology while deliberately withholding private weights, parameter counts, training-data composition, deployment topology, and unreleased benchmark results.", PALE_GRAY)

    add_heading(doc, "1. Introduction", 1)
    add_text(doc, "A contemporary general-purpose model is expected to do far more than continue text. It must synthesize across files, reason over ambiguous requirements, interpret visual artifacts, invoke tools safely, recover from errors, and return results in formats that downstream systems can consume. These demands expose a gap between a single-pass language-model interaction and a durable intelligent system. Quartz addresses that gap by treating the model, its memory behavior, its reasoning budget, its tool interface, and its verification procedures as one integrated architecture.")
    add_text(doc, "The Quartz thesis is intentionally model-first. Its public surface is a family of modes sharing one private-weights identity and a common behavioral contract. Inside that system, the model uses conditional computation and adaptive inference to allocate capacity where it matters. The purpose is not to fragment an interaction into unrelated services; it is to give one coherent model substrate the ability to use distinct internal circuits, memory paths, and inference procedures for distinct kinds of work.")
    add_callout(doc, "Design principle", "Quartz should remain legible as one model system: the user supplies an objective, the system forms and updates an internal working state, and the final response carries a single coherent identity, safety posture, and reliability contract.")
    add_figure(doc, figures[0], "Figure 1", "Quartz system overview")

    add_heading(doc, "2. Design Objectives", 1)
    add_text(doc, "Quartz is organized around six design objectives. First, it must preserve useful evidence across very long inputs without making every token interaction equally expensive. Second, it must expose a large private capacity while keeping per-request computation bounded through conditional expert activation. Third, it must vary test-time effort with task difficulty, ambiguity, and consequence. Fourth, it must work natively with tools and structured interfaces rather than treating them as prompt-only conventions. Fifth, it must support multimodal evidence as part of the same reasoning workspace. Sixth, it must make reliability observable through verification, traceable tool results, schema conformance, and evaluation gates.")
    add_table(doc, ["Objective", "Design response", "Observable evidence"], [
        ["Long-context fidelity", "Hybrid sequence pathways, stable-prefix reuse, segment attribution", "Retrieval fidelity and multi-document synthesis tests"],
        ["High capacity at bounded cost", "Sparse internal expert activation with shared pathways", "Load balance, expert utilization, quality-cost curves"],
        ["Adaptive reasoning", "Difficulty estimation and verifier-guided expansion", "Calibration, improvement per extra inference step"],
        ["Agent reliability", "Schema-bound tools, permission boundaries, observation loops", "Tool success, recovery, and unsafe-action refusal tests"],
        ["Machine usability", "Strict structured decoding and continuation support", "Schema validity, parsing success, deterministic contract tests"],
    ], [1.35, 3.25, 1.9])

    add_heading(doc, "3. Architectural Overview", 1)
    add_text(doc, "Quartz can be read as five layers arranged along a single inference path. The input fabric turns text, code, documents, images, and tool observations into typed model context. The core model fabric transforms that context through a long-sequence backbone and a sparse expert substrate. The inference fabric decides how much internal work a request merits and whether it requires a tool-grounded loop. The verification fabric checks intermediate and final candidates against task constraints. The response fabric emits natural language, structured data, actions, and streams under one output contract.")
    add_text(doc, "This decomposition is not merely diagrammatic. It makes technical responsibilities explicit. The context layer owns provenance and stable reuse. The backbone owns representation learning. The expert layer owns conditional specialization. The inference layer owns deliberation depth. The tool layer owns executable interaction. The output layer owns schema, presentation, and transport. By separating responsibilities while retaining a shared model state, Quartz avoids the common failure mode in which agent behavior becomes an opaque collection of prompt templates.")

    add_heading(doc, "4. Long-Context Representation", 1)
    add_heading(doc, "4.1 Hybrid Sequence Pathway", 2)
    add_text(doc, "Long-context quality is not equivalent to accepting a large number of tokens. A model must preserve local details, carry information over distant spans, and recover source-specific evidence when it is needed. Full attention offers rich token-to-token interaction but has unfavorable growth as sequence length rises. Linear or recurrent-style pathways scale more gently, but can lose the expressive precision required for local code, tables, or cross-reference resolution. Quartz therefore uses a hybrid sequence design: a selective attention path handles high-information neighborhoods and retrieval-sensitive interactions, while a gated linear memory path carries a compact state over extended spans. A learned fusion mechanism combines both paths at each block.")
    add_text(doc, "The practical goal is not to choose one universal attention mechanism. It is to recognize that model context contains different information geometries. A source citation, a function signature, and a table header often require sharp local alignment. A large document history, a user preference, or an evolving plan may benefit more from stable state transport. The hybrid pathway allows Quartz to reserve high-resolution attention for the first class of signals and a more economical memory process for the second. Gated linear attention research provides a useful antecedent: linear-time recurrent formulations can be competitive when gating and hardware-aware execution recover much of the expressiveness lost by simpler linear mechanisms [2].")
    add_figure(doc, figures[1], "Figure 2", "Hybrid long-context sequence core")
    add_heading(doc, "4.2 Attention Residual Transport", 2)
    add_text(doc, "Deep sequence models can lose useful attention-derived information as hidden states pass through many transformations. Quartz uses explicit residual transport around attention-derived features as a design requirement. The residual path is not a substitute for learning; it is a controlled pathway through which salient token relationships can remain recoverable. The mixer learns when to preserve, amplify, or attenuate that signal. This improves the model's ability to reconcile recent precision with longer-running state, especially in contexts that mix short critical spans with large reference material.")
    add_heading(doc, "4.3 Context Assembly, Caching, and Provenance", 2)
    add_text(doc, "Context is assembled as a structured object, not as an undifferentiated prompt string. Stable prefixes such as system instructions, repository indexes, or recurring knowledge packs can be keyed and reused when unchanged. Mutable segments such as a new user turn, a diff, or a tool result are composed around that stable state. A segment index records source identity, position, and summary anchors. This permits the model to revisit relevant spans and allows the verification layer to distinguish model inference from tool-observed evidence.")
    add_figure(doc, figures[3], "Figure 3", "Long-context memory fabric")

    add_heading(doc, "5. Private Sparse Expert Core", 1)
    add_text(doc, "Quartz exposes a single private-weights model identity while internally using sparse, conditionally activated expert circuits. Mixture-of-experts research demonstrates that conditional activation can increase model capacity without requiring every input to traverse every parameter group [3]. For Quartz, the important point is architectural coherence: expert selection occurs inside the model substrate at token or span granularity. It is an internal computation decision that is fused back into shared hidden states, not a change in external identity or a handoff between unrelated products.")
    add_text(doc, "The sparse-expert block begins with shared normalized hidden states. A learned selector computes affinity scores over a large internal expert inventory and selects a small active set under capacity constraints. Selected experts may specialize along emergent dimensions such as formal reasoning, code transformation, multilingual synthesis, planning, or tool-state interpretation. The selector is trained with load-balance and stability objectives so that it does not collapse onto a small number of popular experts. Expert outputs are weighted, merged, and passed through residual pathways that preserve general instruction-following behavior.")
    add_figure(doc, figures[2], "Figure 4", "Stable sparse expert execution")
    add_heading(doc, "5.1 Stability Requirements", 2)
    add_text(doc, "Conditional capacity adds failure modes that dense models do not have. Expert load can become imbalanced; expert selection can become brittle under distribution shift; and an expert can over-specialize to superficial cues. Quartz therefore treats three controls as first-class: auxiliary load regularization, selection diversity monitoring, and capacity-aware fallbacks. During training, utilization statistics and cross-domain validation detect collapse early. During inference, the shared trunk and residual pathways ensure that a degraded expert choice does not erase a model-wide representation. This is particularly important for long-context and agentic tasks, where token distributions can shift sharply as tool observations enter the conversation.")

    add_heading(doc, "6. Adaptive Inference and Test-Time Reasoning", 1)
    add_text(doc, "A single fixed inference budget is poorly matched to heterogeneous work. A short transformation with explicit instructions should be fast. A mathematical proof, a codebase change, or a high-consequence structured action may merit additional planning, candidate generation, checking, and revision. Quartz implements adaptive inference as a continuous control problem. A difficulty estimator reads task features such as scope, ambiguity, number of constraints, context complexity, prior failure signals, and required tool interaction. It then selects a reasoning policy and a bounded compute budget.")
    add_text(doc, "The direct-response policy is appropriate when the model has high confidence and the result is easy to validate. The reasoning policy creates a working plan, derives candidate solutions, and applies focused critique. The tool-grounded policy adds environment interaction: the model proposes schema-valid actions, observes results, updates state, and either continues or terminates. Across all policies, a verifier ranks or rejects candidates based on task-specific criteria. Process supervision is a useful reference point here: step-level feedback can improve reliability on multi-step reasoning tasks because it identifies an error before it reaches a plausible-looking final answer [4].")
    add_figure(doc, figures[4], "Figure 5", "Adaptive inference and verification")
    add_heading(doc, "6.1 Verifier-Guided Candidate Selection", 2)
    add_text(doc, "Verification should be discriminating rather than ceremonial. Quartz verifiers operate at multiple granularities. A syntax verifier can test whether a structured answer conforms to a schema. A code verifier can run compilation or tests in an authorized environment. A factuality verifier can check whether cited evidence is present in the assembled context or returned by a tool. A reasoning verifier can inspect whether a candidate violates explicit constraints or contains internal inconsistencies. The verifier does not need to recreate the answer; it needs a narrow, auditable criterion that separates acceptable and unacceptable candidates. This lets Quartz spend additional compute where a detectable uncertainty remains.")

    add_heading(doc, "7. Tool-Grounded Agentic Execution", 1)
    add_text(doc, "Tool use turns language-model inference into stateful control. That creates power and risk at the same time. Quartz therefore treats tool calls as typed transitions, not prose suggestions. A tool declaration specifies a name, purpose, JSON schema, permission scope, side-effect class, and result schema. The model may only emit calls that satisfy the declared structure. An execution policy then determines whether the call is allowed in the current task and whether it requires confirmation, sandboxing, or a more restrictive mode.")
    add_text(doc, "The core loop is goal, plan, act, observe, verify, and synthesize. Reasoning traces formulate a plan and track progress; actions retrieve or transform external state; observations become typed context; and verification checks whether the goal has advanced or a recovery path is required. This approach follows the broader lesson from reasoning-and-acting research: external actions can reduce hallucination and enable adaptive plans when their results are integrated back into the model state [5]. Quartz adds bounded authority and explicit stopping rules so that an agent does not continue simply because it can.")
    add_figure(doc, figures[5], "Figure 6", "Tool-grounded execution loop")
    add_heading(doc, "7.1 Dynamic Capability Loading", 2)
    add_text(doc, "Agent tasks do not need every tool exposed at every moment. Overexposure increases prompt size, tool-selection ambiguity, and attack surface. Quartz supports dynamic capability loading: a planner may request a relevant tool definition when the task reaches a subgoal that needs it. The definition remains part of the model context for the duration of the applicable workflow, so the model can reason over its complete schema and examples. This keeps the operational surface smaller while preserving composability.")
    add_heading(doc, "7.2 Deterministic and External Computation", 2)
    add_text(doc, "Some classes of work benefit from deterministic computation rather than additional natural-language inference. Arithmetic, parsing, code execution, retrieval, and file operations should be delegated to authorized deterministic tools when the task calls for them. Quartz uses the model to decide what to compute, explain the result, and integrate observed outputs. The tool, not the model, produces the authoritative external result. This division makes failures easier to diagnose and makes the final response more grounded.")

    add_heading(doc, "8. Output Reliability and Developer Interface", 1)
    add_text(doc, "The last stage of generation should not be a mere text stream. Quartz supports a response contract that distinguishes internal working state from user-visible content. The output decoder may target natural language, strict structured data, tool calls, partial continuations, or multi-channel streams. For structured tasks, decoding is constrained by an explicit JSON Schema or equivalent grammar. The decoder masks invalid continuations and the verifier checks the final object before it is emitted. Structured-generation research shows why this matters: schema compliance is measurable, and real-world schemas exercise a broader range of constraints than toy examples [6].")
    add_text(doc, "Partial continuation is useful when a developer needs the model to complete a preexisting prefix, a document scaffold, or an incremental code stream. The prefix is treated as a deliberate part of the response contract, not as a hidden prompt transformation. Streaming also benefits from separation: progress or reasoning-status events can be transported distinctly from final user-facing content and schema-bound artifacts. The final response remains the authoritative result; intermediate progress must not be mistaken for a verified conclusion.")
    add_figure(doc, figures[6], "Figure 7", "Reliable output contract")

    add_heading(doc, "9. Training and Alignment Program", 1)
    add_text(doc, "A credible private model system requires a training program that aligns architectural choices with data, objectives, and evaluation. Quartz training begins with data curation: deduplication, quality filtering, document-structure preservation, code validity checks, multimodal alignment, contamination control, and domain balancing. Foundation training then emphasizes sequence stability, long-context curricula, sparse-expert utilization, and robust instruction representation. Capability shaping adds task families that exercise reasoning, code generation and repair, tool calling, structured output, and multimodal synthesis.")
    add_text(doc, "Post-training should combine preference data with process-aware signals. Preference learning helps calibrate helpfulness, style, and policy adherence. Process signals help identify where a multi-step trajectory deviates from a valid path. Tool tasks should include observation noise, API errors, stale state, and permission failures rather than only clean happy paths. Alignment is complete only when it is paired with adversarial evaluation: prompt injection, deceptive tool descriptions, sensitive-data requests, unsafe action chains, and long-context distraction attacks are all relevant to a system that can act.")
    add_figure(doc, figures[7], "Figure 8", "Training and alignment lifecycle")

    add_heading(doc, "10. Evaluation Methodology", 1)
    add_text(doc, "Quartz evaluation is a release discipline, not a one-time leaderboard event. The benchmark program should use versioned task suites, hidden holdouts, controlled ablations, and failure-trace review. Scores must be joined with compute, latency, schema validity, tool success, and safety measures. A model that improves a headline reasoning score while silently increasing unsafe action rates or degrading long-context attribution has not improved in the way Quartz is designed to improve.")
    add_figure(doc, figures[8], "Figure 9", "Evaluation and release evidence")
    add_table(doc, ["Area", "Representative task", "Primary measure", "Failure evidence"], [
        ["Reasoning", "Constraint satisfaction, proofs, planning", "Accuracy with calibrated confidence", "Invalid step, contradiction, unsupported conclusion"],
        ["Software engineering", "Issue resolution in real repositories", "Test pass rate and patch validity", "Regression, incomplete edit, fabricated file state"],
        ["Long context", "Multi-document retrieval and synthesis", "Evidence recall, attribution, distractor resistance", "Lost source, wrong binding, stale cache use"],
        ["Structured output", "Nested schema extraction and generation", "Parse rate and schema conformance", "Invalid object, omitted requirement, extra field"],
        ["Tools and agents", "Multi-step environment tasks", "Goal completion under bounded authority", "Unsafe call, loop, ignored observation, bad recovery"],
        ["Multimodal", "Document, screenshot, and visual reasoning", "Grounded answer quality", "Visual hallucination, missed region, unsupported claim"],
    ], [1.15, 1.85, 1.75, 1.75])
    add_heading(doc, "10.1 No Unreleased Claims", 2)
    add_text(doc, "This paper intentionally does not publish parameter counts, training-token counts, benchmark ranks, or comparisons to other named systems. Those claims require a separate report with fixed model versions, disclosed evaluation protocol, confidence intervals where appropriate, and reproducible artifacts or clearly stated access limitations. The purpose here is to make Quartz's design logic inspectable before a performance report is issued.")

    add_heading(doc, "11. Security, Privacy, and Safety", 1)
    add_text(doc, "Quartz is designed for private weights and controlled operations. That architecture does not itself guarantee security; the deployment must enforce it. Input isolation, tenant boundaries, access logging, secrets handling, data retention rules, and tool credentials are deployment responsibilities. At the model-system layer, Quartz contributes typed context, provenance tags, declared tool permissions, risk-aware inference policies, and stop rules. These controls make safety behavior observable and testable rather than relying on an instruction string to carry the full burden.")
    add_table(doc, ["Risk", "Quartz control", "Verification expectation"], [
        ["Prompt injection in external content", "Trust labels, source-aware context, restricted tool activation", "Injected instructions cannot broaden tool authority"],
        ["Unsafe or irreversible actions", "Permission classes, confirmation gates, sandbox defaults", "High-impact calls are blocked or escalated"],
        ["Hallucinated external state", "Tool-result binding and provenance checks", "Claims about tools cite returned observations"],
        ["Long-context distraction", "Segment anchors, targeted retrieval, evidence verification", "Critical constraints survive distractor-heavy inputs"],
        ["Structured interface drift", "Strict schemas and contract tests", "Invalid outputs are blocked"],
    ], [1.55, 2.75, 2.05])

    add_heading(doc, "12. Limitations and Open Research Questions", 1)
    add_text(doc, "The architecture described here is ambitious and has real tradeoffs. Hybrid sequence paths introduce implementation complexity and require careful calibration of which interactions deserve high-resolution attention. Sparse experts improve conditional capacity but create expert-selection stability and systems challenges. Adaptive inference can increase reliability, yet poorly calibrated expansion wastes compute or creates a false sense of certainty. Tool use grounds some tasks while increasing the need for environment isolation and permission design. Verification is only as good as the criterion it can test; open-ended questions may remain difficult to score without overfitting to superficial form.")
    add_text(doc, "Several questions remain open. How should long-context cache reuse be validated when source documents change subtly? How can expert selection remain interpretable without revealing sensitive internal weights? When should a verifier trigger another reasoning pass instead of returning uncertainty to the user? How should a multimodal model preserve exact visual grounding across long videos or dense technical diagrams? And how can evaluation distinguish genuine agentic competence from benchmark-specific tool patterns? Quartz treats these as research questions, not solved marketing claims.")

    add_heading(doc, "13. Conclusion", 1)
    add_text(doc, "Quartz is designed as a coherent private model system for work that exceeds a single text completion. Its technical thesis combines hybrid long-context representation, stable sparse expert computation, adaptive inference, verifier-guided selection, typed tool execution, and schema-constrained output. These mechanisms are mutually reinforcing: memory makes evidence available; experts provide conditional capacity; adaptive inference spends time where it matters; tools connect the model to authoritative state; verification protects the response boundary; and a unified output contract makes the system usable by people and software alike.")
    add_text(doc, "The next legitimate step is measurement. A separate technical report should publish fixed-version evaluation evidence, operational envelopes, and limitations once they are ready for review. Until then, the contribution of this paper is an explicit architecture: what Quartz is intended to be, how its major subsystems fit together, and what a rigorous release standard should demand.")

    add_heading(doc, "References", 1)
    refs = [
        "A. Vaswani et al. Attention Is All You Need. NeurIPS, 2017. https://arxiv.org/abs/1706.03762",
        "S. Yang et al. Gated Linear Attention Transformers with Hardware-Efficient Training. arXiv:2312.06635, 2023. https://arxiv.org/abs/2312.06635",
        "W. Fedus, B. Zoph, and N. Shazeer. Switch Transformers: Scaling to Trillion Parameter Models with Simple and Efficient Sparsity. JMLR, 2022. https://arxiv.org/abs/2101.03961",
        "H. Lightman et al. Let's Verify Step by Step. arXiv:2305.20050, 2023. https://arxiv.org/abs/2305.20050",
        "S. Yao et al. ReAct: Synergizing Reasoning and Acting in Language Models. ICLR, 2023. https://arxiv.org/abs/2210.03629",
        "S. Geng et al. Generating Structured Outputs from Language Models: Benchmark and Studies. arXiv:2501.10868, 2025. https://arxiv.org/abs/2501.10868",
        "W. X. Wang et al. LongBench: A Bilingual, Multitask Benchmark for Long Context Understanding. arXiv:2308.14508, 2023. https://arxiv.org/abs/2308.14508",
        "N. F. Liu et al. Lost in the Middle: How Language Models Use Long Contexts. TACL, 2024. https://arxiv.org/abs/2307.03172",
        "C. B. Mann et al. SWE-bench: Can Language Models Resolve Real-World GitHub Issues? ICLR, 2024. https://arxiv.org/abs/2310.06770",
        "X. Zhou et al. WebArena: A Realistic Web Environment for Building Autonomous Agents. ICLR, 2024. https://arxiv.org/abs/2307.13854",
        "X. Wang et al. Self-Consistency Improves Chain of Thought Reasoning in Language Models. ICLR, 2023. https://arxiv.org/abs/2203.11171",
        "V. Venktesh, M. Rathee, and A. Anand. Trust but Verify! A Survey on Verification Design for Test-time Scaling. arXiv:2508.16665, 2025. https://arxiv.org/abs/2508.16665",
    ]
    for index, ref in enumerate(refs, 1):
        add_reference(doc, index, ref)

    add_heading(doc, "Appendix A. Quartz Operating Modes", 1)
    add_text(doc, "Quartz modes are operational compute profiles, not separate public identities. They maintain the same private-weights model system, tool contract, and safety policy while differing in the default inference budget and latency envelope. Product packaging and access policy are intentionally outside the scope of this technical paper.")
    add_table(doc, ["Mode", "Default posture", "Suitable work", "Escalation behavior"], [
        ["Quartz Lite", "Latency-aware direct synthesis with bounded checking", "Chat, extraction, small edits, routine transformations", "Escalates only for explicit tool requirements or high uncertainty"],
        ["Quartz", "Balanced reasoning and verification", "General knowledge work, coding assistance, document analysis", "Uses a reasoning loop when complexity and consequence rise"],
        ["Quartz Hyper", "Maximum permitted deliberation and verification", "Long-horizon engineering, complex analysis, tool-grounded workflows", "Expands planning, candidates, and verification within safety limits"],
    ], [1.25, 2.0, 1.75, 1.5])

    add_heading(doc, "Appendix B. Diagram Reading Notes", 1)
    add_text(doc, "The diagrams are conceptual. Arrows indicate the primary direction of information or control flow. They do not imply a one-to-one implementation module, public endpoint, parameter group, or deployment boundary. The sparse-expert diagram is especially important: it represents internal conditional circuits within one model system. The agent loop likewise shows a bounded control cycle, not unrestricted autonomy.")

    doc.core_properties.title = "Quartz: A Private-Weights, Long-Context Foundation Model"
    doc.core_properties.author = "Shubhankar Kahali; Kai Keskitalo; Zuzanna Kowalczyk"
    doc.core_properties.subject = "Trumbo Quartz technical architecture paper"
    doc.core_properties.comments = "Technical research manuscript."
    doc.save(OUT)


if __name__ == "__main__":
    build_doc()
    print(OUT)
